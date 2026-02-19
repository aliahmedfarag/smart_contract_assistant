import os
import PyPDF2
from docx import Document
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangChainDocument


class DocumentIngestion:
    
    def __init__(self, chunk_size=800, chunk_overlap=150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        self.vectorstore = None

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    def extract_text_from_docx(self, docx_path: str) -> str:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def process_document(self, file_path: str) -> List[LangChainDocument]:
        if file_path.endswith('.pdf'):
            text = self.extract_text_from_pdf(file_path)
        elif file_path.endswith(('.docx', '.doc')):
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file type")

        chunks = self.text_splitter.split_text(text)

        return [
            LangChainDocument(
                page_content=chunk,
                metadata={"source": os.path.basename(file_path), "chunk_id": i}
            )
            for i, chunk in enumerate(chunks)
        ]

    def ingest_document(self, file_path: str, save_name: str = "default") -> FAISS:
        documents = self.process_document(file_path)
        
        self.vectorstore = FAISS.from_documents(documents, self.embeddings)
        
        os.makedirs("vectorstore", exist_ok=True)
        self.vectorstore.save_local(f"./vectorstore/{save_name}")
        
        return self.vectorstore

    def load_vectorstore(self, name: str = "default") -> FAISS:
        path = f"./vectorstore/{name}"
        if not os.path.exists(path):
            return None
        return FAISS.load_local(path, self.embeddings, allow_dangerous_deserialization=True)